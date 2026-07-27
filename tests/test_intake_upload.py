"""Upload-intake tests (F01 — third intake mode: uploaded documents).

Covers the extraction layer (magic-byte detection, per-type happy path, oversize,
wrong type, empty-text PDF, encrypted PDF) and the router-level tenancy guarantee
that an uploaded document lands under the CALLER's organization only.

No network, no live DB — fixtures are built in-process and the Supabase layer is
patched with capturing fakes.
"""

import io
import time
import zipfile
from unittest.mock import AsyncMock, patch
from uuid import uuid4

import fitz  # PyMuPDF — build tiny PDFs in-memory
import jwt as pyjwt
import pytest
from docx import Document
from httpx import ASGITransport, AsyncClient

from app.config import settings
from app.main import app
from app.services.intake.extract import (
    MAX_UPLOAD_BYTES,
    ExtractionError,
    _detect_upload_kind,
    extract_from_upload,
)

# A privacy-notice-shaped body long enough to decompose into real clauses.
SAMPLE_TEXT = (
    "Privacy Notice\n\n"
    "We collect your name, email address, and browsing data when you use our "
    "services. We use cookies and tracking pixels for analytics.\n\n"
    "We share your personal information with third-party service providers. "
    "You have the right to access, delete, and correct your personal data. "
    "We retain your data for as long as necessary to provide the service."
)


# ── Fixture builders ─────────────────────────────────────────

def _pdf_bytes(text: str) -> bytes:
    doc = fitz.open()
    page = doc.new_page()
    if text:
        page.insert_text((72, 72), text, fontsize=11)
    data = doc.tobytes()
    doc.close()
    return data


def _docx_bytes(paragraphs: list[str]) -> bytes:
    doc = Document()
    doc.add_heading("Privacy Notice", level=1)
    for p in paragraphs:
        doc.add_paragraph(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _png_bytes() -> bytes:
    # 1x1 PNG — a valid non-document binary to exercise wrong-type rejection.
    return bytes.fromhex(
        "89504e470d0a1a0a0000000d49484452000000010000000108060000001f15c489"
        "0000000a49444154789c6360000002000154a24f6f0000000049454e44ae426082"
    )


# ── Magic-byte detection ─────────────────────────────────────

def test_detect_pdf_by_magic():
    assert _detect_upload_kind(_pdf_bytes(SAMPLE_TEXT)) == "pdf"


def test_detect_docx_by_magic():
    assert _detect_upload_kind(_docx_bytes([SAMPLE_TEXT])) == "docx"


def test_detect_txt_by_magic():
    assert _detect_upload_kind(SAMPLE_TEXT.encode()) == "txt"


def test_detect_ignores_extension_uses_bytes():
    # PDF bytes must be recognized as PDF even without any filename hint.
    assert _detect_upload_kind(_pdf_bytes(SAMPLE_TEXT)) == "pdf"


def test_detect_png_rejected():
    with pytest.raises(ExtractionError, match="Unsupported file type"):
        _detect_upload_kind(_png_bytes())


def test_detect_non_word_zip_rejected():
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as z:
        z.writestr("hello.txt", "not a word doc")
    with pytest.raises(ExtractionError, match="ZIP or Office file"):
        _detect_upload_kind(buf.getvalue())


# ── Happy path per type ──────────────────────────────────────

def test_extract_pdf_happy():
    text, chash, kind, mime = extract_from_upload(_pdf_bytes(SAMPLE_TEXT), "notice.pdf")
    assert kind == "pdf"
    assert mime == "application/pdf"
    assert "privacy" in text.lower()
    assert len(chash) == 64


def test_extract_docx_happy():
    body = ["We collect your email address and browsing data.",
            "You have the right to delete your personal data."]
    text, chash, kind, mime = extract_from_upload(_docx_bytes(body), "notice.docx")
    assert kind == "docx"
    assert mime.endswith("wordprocessingml.document")
    assert "# Privacy Notice" in text          # heading mapped to markdown
    assert "right to delete" in text
    assert len(chash) == 64


def test_extract_txt_happy():
    text, chash, kind, mime = extract_from_upload(SAMPLE_TEXT.encode(), "notice.txt")
    assert kind == "txt"
    assert mime == "text/plain"
    assert text.startswith("Privacy Notice")


def test_extract_txt_strips_bom():
    data = ("﻿" + SAMPLE_TEXT).encode("utf-8")
    text, _, kind, _ = extract_from_upload(data, "notice.txt")
    assert kind == "txt"
    assert not text.startswith("﻿")


# ── Oversize ─────────────────────────────────────────────────

def test_oversize_rejected():
    big = b"%PDF-" + b"0" * (MAX_UPLOAD_BYTES + 1)
    with pytest.raises(ExtractionError, match="too large"):
        extract_from_upload(big, "big.pdf")


def test_empty_upload_rejected():
    with pytest.raises(ExtractionError, match="empty"):
        extract_from_upload(b"", "nothing.pdf")


# ── Wrong type ───────────────────────────────────────────────

def test_wrong_type_rejected():
    with pytest.raises(ExtractionError, match="Unsupported file type"):
        extract_from_upload(_png_bytes(), "logo.png")


# ── Empty-text PDF (scanned / no text layer) ─────────────────

def test_empty_text_pdf_rejected():
    # A PDF page with no text layer → honest "no readable text" error, not a crash.
    with pytest.raises(ExtractionError, match="readable text"):
        extract_from_upload(_pdf_bytes(""), "scanned.pdf")


# ── Encrypted PDF ────────────────────────────────────────────

def test_encrypted_pdf_rejected():
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), SAMPLE_TEXT, fontsize=11)
    encrypted = doc.tobytes(
        encryption=fitz.PDF_ENCRYPT_AES_256,
        owner_pw="owner",
        user_pw="secret",  # user password required → unreadable
    )
    doc.close()
    with pytest.raises(ExtractionError, match="password-protected"):
        extract_from_upload(encrypted, "locked.pdf")


# ── Tenancy: upload lands under the caller's org only ────────

_MINE = str(uuid4())
_OTHER = str(uuid4())


def _token(role="customer", org=_MINE):
    now = int(time.time())
    return pyjwt.encode(
        {"sub": "u", "aud": "authenticated", "iat": now - 60, "exp": now + 3600,
         "app_role": role, "organization_id": org},
        settings.supabase_jwt_secret, algorithm="HS256",
    )


class _Resp:
    def __init__(self, data, status_code=200):
        self._data, self.status_code, self.text = data, status_code, ""

    def json(self):
        return self._data


@pytest.mark.anyio
async def test_upload_lands_under_caller_org_only():
    """A customer uploading a document — even if they try to pass a foreign
    organization_id — persists the notice under THEIR org, never another tenant."""
    import app.routers.assessments as A

    captured: dict = {}

    async def _fake_post(table, payload):
        if table == "privacy_notice":
            captured["notice"] = payload
        return _Resp([], 201)

    async def _fake_get(table, *, select="*", filters="", limit=1000, count=False):
        return _Resp([])

    async def _fake_score(*a, **k):
        return {"summary": {"overall_intelligence": 42.0}}

    docx = _docx_bytes(["We collect your email. You may delete your data. "
                        "We share data with providers and retain it as needed."])

    transport = ASGITransport(app=app)
    with patch.object(A, "supabase_rest_post", _fake_post), \
         patch.object(A, "supabase_rest_get", _fake_get), \
         patch("app.services.live_scoring.score_and_persist", AsyncMock(side_effect=_fake_score)), \
         patch("app.services.llm.get_llm_client", side_effect=RuntimeError("no llm in test")):
        async with AsyncClient(transport=transport, base_url="http://test") as c:
            r = await c.post(
                "/assessments/",
                files={"file": ("notice.docx", docx,
                                "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                data={"organization_id": _OTHER},  # attempted cross-tenant write
                headers={"Authorization": f"Bearer {_token('customer', _MINE)}"},
            )

    assert r.status_code == 201, r.text
    body = r.json()
    assert body["intake_method"] == "upload"
    assert body["ssrf_protected"] is False        # uploads are never verified-source
    assert body["organization_id"] == _MINE       # NOT _OTHER
    assert captured["notice"]["organization_id"] == _MINE
    assert captured["notice"]["intake_method"] == "upload"
    assert captured["notice"]["upload_filename"] == "notice.docx"
    assert captured["notice"]["upload_mime"].endswith("wordprocessingml.document")
    assert len(captured["notice"]["upload_file_hash"]) == 64
