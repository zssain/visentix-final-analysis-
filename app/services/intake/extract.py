"""Text extraction from URL, PDF upload, or raw text.

All extraction produces CLEAN PLAIN TEXT ready for decomposition. Output is
structured with markdown headings (# / ## / ###) and blank-line-separated
paragraphs so that decompose._split_sections picks them up correctly.

Defenses:
- SSRF: every redirect hop is re-validated (no TOCTOU via 302)
- Size limits on response body, PDF, and text
- MIME validation on URL fetches (not just uploads)
- HTML cleaning strips scripts/styles/nav/chrome → body text only
- Privacy-policy signal check (advisory, not blocking)
"""

from __future__ import annotations

import hashlib
import io
import re
import zipfile
from urllib.parse import urljoin

import httpx
from bs4 import BeautifulSoup

from app.services.intake.ssrf import (
    FETCH_TIMEOUT_SECONDS,
    MAX_RESPONSE_BYTES,
    SSRFError,
    validate_url,
)

# ── Constants ────────────────────────────────────────────────

MAX_TEXT_LENGTH = 500_000  # 500K chars
MAX_PDF_BYTES = 10 * 1024 * 1024  # 10 MB
# Uploaded documents (PDF / DOCX / TXT) share one size ceiling. Matches the
# F01 spec's ≤10 MB intake limit; NOT a separate value.
MAX_UPLOAD_BYTES = MAX_PDF_BYTES
MAX_REDIRECTS = 3
MIN_POLICY_TEXT_CHARS = 300

ALLOWED_MIME_TYPES = {
    "application/pdf",
    "text/html",
    "text/plain",
    "application/xhtml+xml",
}

# Canonical MIME per detected upload kind — recorded as honest provenance on the
# source record. Detection is by MAGIC BYTES (below), never the client's
# Content-Type header or the filename extension (both trivially spoofable).
UPLOAD_MIME = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "txt": "text/plain",
}

USER_AGENT = (
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
    "AppleWebKit/537.36 (KHTML, like Gecko) "
    "Chrome/130.0.0.0 Safari/537.36"
)

# Tags whose content is never privacy-notice text
_STRIP_TAGS = (
    "script", "style", "noscript", "svg", "img", "nav", "header",
    "footer", "iframe", "form", "button", "aside", "template",
)

# Substrings that signal the page is likely a privacy policy
_PRIVACY_SIGNALS = (
    "privacy", "personal information", "data", "cookies", "collect",
)


class ExtractionError(ValueError):
    pass


# ── Public helpers ───────────────────────────────────────────

def looks_like_privacy_policy(text: str) -> bool:
    """Heuristic check: does the text look like a privacy policy?

    Counts how many of the 5 privacy signals appear as substrings.
    Returns True if >= 3 are present. Used for advisory warnings,
    not hard blocking (spec principle: don't refuse unusual pages).
    """
    lower = text.lower()
    hits = sum(1 for signal in _PRIVACY_SIGNALS if signal in lower)
    return hits >= 3


# ── SSRF-safe fetch with redirect re-validation ─────────────

async def _fetch_ssrf_safe(url: str) -> httpx.Response:
    """Fetch a URL, re-validating SSRF on every redirect hop.

    - follow_redirects=False so WE control the redirect chain
    - Each hop runs validate_url() (DNS resolve + IP block check)
    - Max MAX_REDIRECTS hops, then error
    """
    current = validate_url(url)  # raises SSRFError if blocked

    async with httpx.AsyncClient(
        timeout=FETCH_TIMEOUT_SECONDS,
        follow_redirects=False,
        headers={
            "User-Agent": USER_AGENT,
            "Accept-Language": "en-US,en;q=0.9",
        },
    ) as client:
        for _ in range(MAX_REDIRECTS + 1):
            response = await client.get(current)

            if response.status_code in (301, 302, 303, 307, 308):
                location = response.headers.get("location")
                if not location:
                    raise ExtractionError(
                        f"Redirect {response.status_code} without Location header"
                    )
                # Resolve relative URLs and re-validate for SSRF
                current = validate_url(urljoin(current, location))
                continue

            response.raise_for_status()
            return response

    raise ExtractionError(f"Too many redirects (max {MAX_REDIRECTS})")


# ── HTML → structured plain text ─────────────────────────────

def _html_to_text(html: str) -> str:
    """Convert HTML to structured plain text with markdown headings.

    Output format matches what decompose._split_sections expects:
    - # / ## / ### headings from <h1>–<h6>
    - Blank-line-separated paragraphs from <p>, <li>, etc.
    - All script/style/nav/footer chrome stripped
    """
    soup = BeautifulSoup(html, "lxml")

    # Strip non-content tags
    for tag in soup(_STRIP_TAGS):
        tag.decompose()

    block_tags = [
        "h1", "h2", "h3", "h4", "h5", "h6",
        "p", "li", "blockquote", "pre", "td", "dt", "dd",
    ]
    heading_prefix = {
        "h1": "# ", "h2": "## ", "h3": "### ",
        "h4": "### ", "h5": "### ", "h6": "### ",
    }

    blocks: list[str] = []
    seen_elements: set[int] = set()  # track element ids to avoid duplicates

    for el in soup.find_all(block_tags):
        el_id = id(el)
        if el_id in seen_elements:
            continue

        # Skip a <p> (or similar) nested inside an already-captured container
        # e.g., <li><p>text</p></li> — the <li> already captured this text
        if el.name in ("p", "dd") and el.parent and el.parent.name in ("li", "td", "blockquote"):
            parent_id = id(el.parent)
            if parent_id in seen_elements:
                continue

        text = el.get_text(" ", strip=True)
        text = re.sub(r"[ \t]+", " ", text)

        if not text:
            continue

        seen_elements.add(el_id)

        prefix = heading_prefix.get(el.name, "")
        blocks.append(f"{prefix}{text}")

    if blocks:
        return "\n\n".join(blocks).strip()

    # Fallback: no block elements found (malformed HTML)
    fallback = soup.get_text(separator="\n", strip=True)
    fallback = re.sub(r"[ \t]+", " ", fallback)
    fallback = re.sub(r"\n{3,}", "\n\n", fallback)
    return fallback.strip()


# ── Public extraction functions ──────────────────────────────

async def extract_from_url(url: str) -> tuple[str, str]:
    """Fetch URL content and extract clean text. Returns (text, content_hash).

    SSRF-safe: re-validates on every redirect hop.
    HTML cleaned: strips scripts/styles/nav, produces structured markdown.
    MIME validated: rejects non-allowed content types.
    """
    response = await _fetch_ssrf_safe(url)

    # Size check
    content_length = len(response.content)
    if content_length > MAX_RESPONSE_BYTES:
        raise ExtractionError(
            f"Response too large: {content_length} bytes (max {MAX_RESPONSE_BYTES})"
        )

    # MIME check
    raw_ct = response.headers.get("content-type", "")
    content_type = raw_ct.split(";")[0].strip().lower() if raw_ct else ""

    if content_type and content_type not in ALLOWED_MIME_TYPES:
        raise ExtractionError(
            f"Content type '{content_type}' not allowed. "
            f"Accepted: {', '.join(sorted(ALLOWED_MIME_TYPES))}"
        )

    # Extract text based on content type
    if content_type == "application/pdf":
        text = _extract_pdf_text(response.content)
    elif content_type == "text/plain":
        text = response.text
    else:
        # text/html, application/xhtml+xml, or missing header (treat as HTML)
        text = _html_to_text(response.text)

    text = text.strip()

    # Minimum viable content check
    if len(text) < MIN_POLICY_TEXT_CHARS:
        raise ExtractionError(
            f"Page yielded only {len(text)} characters of visible text "
            f"(minimum {MIN_POLICY_TEXT_CHARS}). Check that the URL points "
            f"directly at the privacy policy page."
        )

    # Truncate
    if len(text) > MAX_TEXT_LENGTH:
        text = text[:MAX_TEXT_LENGTH]

    content_hash = hashlib.sha256(text.encode()).hexdigest()
    return text, content_hash


def extract_from_pdf(pdf_bytes: bytes, filename: str = "") -> tuple[str, str]:
    """Extract text from PDF bytes. Returns (text, content_hash).

    Uses PyMuPDF (fitz) — no shell-outs.
    """
    if len(pdf_bytes) > MAX_PDF_BYTES:
        raise ExtractionError(
            f"PDF too large: {len(pdf_bytes)} bytes (max {MAX_PDF_BYTES})"
        )

    text = _extract_pdf_text(pdf_bytes)

    if len(text) > MAX_TEXT_LENGTH:
        text = text[:MAX_TEXT_LENGTH]

    content_hash = hashlib.sha256(text.encode()).hexdigest()
    return text, content_hash


def extract_from_text(raw_text: str) -> tuple[str, str]:
    """Accept raw text input. Returns (text, content_hash)."""
    if len(raw_text) > MAX_TEXT_LENGTH:
        raise ExtractionError(
            f"Text too long: {len(raw_text)} chars (max {MAX_TEXT_LENGTH})"
        )

    text = raw_text.strip()
    if not text:
        raise ExtractionError("Empty text")

    content_hash = hashlib.sha256(text.encode()).hexdigest()
    return text, content_hash


def _extract_pdf_text(pdf_bytes: bytes) -> str:
    """Extract text from PDF using PyMuPDF (no shell-outs).

    Password-protected PDFs raise a plain-language ExtractionError (we cannot
    read them, and saying so honestly beats a cryptic failure). Scanned PDFs
    with no text layer return an empty string; the caller decides how to
    surface that (OCR is a deliberate future item, not in this pass).
    """
    import fitz  # PyMuPDF

    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    except Exception as e:
        raise ExtractionError(f"Could not open the PDF (it may be corrupted): {e}")

    try:
        if doc.needs_pass:  # password-protected → text is unreadable
            raise ExtractionError(
                "This PDF is password-protected, so its text can't be read. "
                "Remove the password and try again, or paste the text directly."
            )
        pages = [page.get_text() for page in doc]
    finally:
        doc.close()

    return "\n\n".join(pages)


# ── Uploaded-document intake (PDF / DOCX / TXT) ──────────────

def _detect_upload_kind(data: bytes) -> str:
    """Identify an uploaded file by MAGIC BYTES, not its extension.

    Returns 'pdf' | 'docx' | 'txt'. Raises ExtractionError for anything else so
    a renamed executable or image can never reach an extractor.
    """
    head = data[:2048]

    # PDF — '%PDF-' signature (usually at byte 0; scan the header to tolerate a
    # small leading BOM/whitespace some generators emit).
    if b"%PDF-" in head[:1024]:
        return "pdf"

    # DOCX — an OOXML file is a ZIP ('PK\x03\x04'). Distinguish a Word document
    # from an arbitrary zip by the presence of the Word body part.
    if data[:2] == b"PK":
        try:
            with zipfile.ZipFile(io.BytesIO(data)) as z:
                names = z.namelist()
        except zipfile.BadZipFile:
            names = []
        if any(n == "word/document.xml" or n.startswith("word/") for n in names):
            return "docx"
        raise ExtractionError(
            "This looks like a ZIP or Office file we can't read as a document. "
            "Please upload a PDF, Word (.docx), or plain-text (.txt) file."
        )

    # TXT — no binary signature. Accept only if it has no NUL bytes (a strong
    # binary marker) and decodes as UTF-8.
    if b"\x00" not in head:
        try:
            data.decode("utf-8-sig")
            return "txt"
        except UnicodeDecodeError:
            pass

    raise ExtractionError(
        "Unsupported file type. Please upload a PDF, Word (.docx), or "
        "plain-text (.txt) file."
    )


def _extract_docx_text(data: bytes) -> str:
    """Extract text from a .docx, mapping Word heading styles to markdown
    headings so decompose()._split_sections finds sections exactly as it does
    for URL/HTML intake — one shared decomposition path, no parallel logic.
    """
    try:
        from docx import Document
    except ImportError:  # pragma: no cover - dependency guard
        raise ExtractionError(
            "Word document support isn't available on the server right now. "
            "Please paste the notice text or upload a PDF."
        )

    try:
        doc = Document(io.BytesIO(data))
    except Exception as e:
        raise ExtractionError(
            f"Could not open the Word document (it may be corrupted): {e}"
        )

    blocks: list[str] = []
    for para in doc.paragraphs:
        text = " ".join(para.text.split())
        if not text:
            continue
        style = (para.style.name if para.style else "") or ""
        if style == "Title" or style.startswith("Heading 1"):
            blocks.append(f"# {text}")
        elif style.startswith("Heading 2"):
            blocks.append(f"## {text}")
        elif style.startswith("Heading"):  # Heading 3+
            blocks.append(f"### {text}")
        else:
            blocks.append(text)

    # Many notices carry rights/retention detail in tables — keep that text.
    for table in doc.tables:
        for row in table.rows:
            cells = [" ".join(c.text.split()) for c in row.cells]
            line = " · ".join(cell for cell in cells if cell)
            if line:
                blocks.append(line)

    return "\n\n".join(blocks)


def extract_from_upload(data: bytes, filename: str = "") -> tuple[str, str, str, str]:
    """Validate + extract an uploaded document. Returns (text, content_hash,
    kind, mime) where kind ∈ {'pdf','docx','txt'} and content_hash is the hash
    of the EXTRACTED TEXT (same semantics as URL/paste intake, so corpus
    change-detection stays consistent).

    Enforces the shared upload size limit and rejects unreadable, encrypted, or
    zero-text files with plain-language errors (Rule 9 — no jargon). Produces the
    same clean-text shape the other intake modes feed into decompose().
    """
    if not data:
        raise ExtractionError("The uploaded file is empty.")
    if len(data) > MAX_UPLOAD_BYTES:
        limit_mb = MAX_UPLOAD_BYTES // (1024 * 1024)
        raise ExtractionError(
            f"File is too large ({len(data) / (1024 * 1024):.1f} MB). "
            f"The maximum upload size is {limit_mb} MB."
        )

    kind = _detect_upload_kind(data)

    if kind == "pdf":
        text = _extract_pdf_text(data)
        empty_msg = (
            "We couldn't find any readable text in this PDF — it looks like a "
            "scanned image rather than a text document. Upload a text-based PDF, "
            "or paste the notice text directly."
        )
    elif kind == "docx":
        text = _extract_docx_text(data)
        empty_msg = (
            "We couldn't find any readable text in this Word document. "
            "Paste the notice text directly, or try a different file."
        )
    else:  # txt
        text = data.decode("utf-8-sig", errors="replace")
        empty_msg = "The uploaded text file has no readable content."

    text = text.strip()
    if not text:
        raise ExtractionError(empty_msg)

    if len(text) > MAX_TEXT_LENGTH:
        text = text[:MAX_TEXT_LENGTH]

    content_hash = hashlib.sha256(text.encode()).hexdigest()
    return text, content_hash, kind, UPLOAD_MIME[kind]
